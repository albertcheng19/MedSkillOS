#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Word document conversion module
Supports Word (.docx) conversion to PDF and text extraction"""

import os
import tempfile
from pathlib import Path
from typing import Optional


class WordConverter:
    """Word document converter"""

    def __init__(self):
        self.supported_formats = [".docx", ".doc"]

    def is_word_file(self, file_path: str) -> bool:
        """Check if it is a Word document"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_formats

    def extract_text(self, docx_path: str) -> str:
        """Extract plain text from Word document

        Args:
            docx_path: Word document path

        Returns:
            Extracted plain text content"""
        try:
            from docx import Document

            doc = Document(docx_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            return "\n\n".join(paragraphs)

        except ImportError:
            raise ImportError("Need to install python-docx library: pip install python-docx")

    def convert_to_pdf(
        self,
        docx_path: str,
        output_path: Optional[str] = None,
        timeout: int = 60,
    ) -> str:
        """Convert Word document to PDF

        Args:
            docx_path: Word document path
            output_path: PDF output path (default: same name.pdf in the same directory)
            timeout: timeout (seconds)

        Returns:
            Generated PDF file path"""
        try:
            from docx2pdf import convert

            docx_path = Path(docx_path)

            if output_path is None:
                output_path = docx_path.with_suffix(".pdf")
            else:
                output_path = Path(output_path)

            if not output_path.parent.exists():
                output_path.parent.mkdir(parents=True)

            convert(str(docx_path), str(output_path))

            return str(output_path)

        except ImportError:
            raise ImportError("You need to install the docx2pdf library: pip install docx2pdf")
        except Exception as e:
            raise RuntimeError(f"PDFConversion failed: {str(e)}")

    def convert_to_text(self, docx_path: str) -> str:
        """Convert Word document to text (also generate PDF)

        Args:
            docx_path: Word document path

        Returns:
            Extracted plain text content"""
        if not self.is_word_file(docx_path):
            raise ValueError(f"Unsupported file format: {docx_path}")

        return self.extract_text(docx_path)

    def process_document(
        self,
        input_path: str,
        output_dir: Optional[str] = None,
        generate_pdf: bool = True,
    ) -> dict:
        """Process Word documents: extract text and optionally generate PDF

        Args:
            input_path: input file path
            output_dir: output directory
            generate_pdf: whether to generate PDF

        Returns:
            Process result dictionary"""
        input_path = Path(input_path)

        if not input_path.exists():
            raise FileNotFoundError(f"File does not exist: {input_path}")

        if not self.is_word_file(str(input_path)):
            raise ValueError(f"noWorddocument: {input_path}")

        result = {
            "input_path": str(input_path),
            "text": None,
            "pdf_path": None,
        }

        text = self.extract_text(str(input_path))
        result["text"] = text

        if generate_pdf:
            if output_dir:
                output_dir = Path(output_dir)
            else:
                output_dir = input_path.parent

            pdf_path = output_dir / f"{input_path.stem}.pdf"
            result["pdf_path"] = self.convert_to_pdf(str(input_path), str(pdf_path))

        return result


def docx_to_text(docx_path: str) -> str:
    """Shortcut function: Word to plain text"""
    converter = WordConverter()
    return converter.convert_to_text(docx_path)


def docx_to_pdf(docx_path: str, output_path: Optional[str] = None) -> str:
    """Shortcut function: Word to PDF"""
    converter = WordConverter()
    return converter.convert_to_pdf(docx_path, output_path)
