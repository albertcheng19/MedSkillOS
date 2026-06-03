#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""format adjustment module
Adjust the paper format according to the journal or school requirements"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Any


@dataclass
class FormatConfig:
    """Format configuration"""

    name: str = "Default"
    version: str = "1.0"

    font: Dict[str, Any] = field(
        default_factory=lambda: {
            "body": "Times New Roman",
            "body_size": 12,
            "title": "Times New Roman",
            "title_size": 16,
            "caption": "Times New Roman",
            "caption_size": 10,
            "header": "Times New Roman",
            "footer": "Times New Roman",
        }
    )

    spacing: Dict[str, Any] = field(
        default_factory=lambda: {
            "line_space": "single",
            "paragraph_space": 6,
            "indent": 0.5,
        }
    )

    margins: Dict[str, float] = field(
        default_factory=lambda: {
            "top": 2.54,
            "bottom": 2.54,
            "left": 2.54,
            "right": 2.54,
        }
    )

    references: Dict[str, str] = field(
        default_factory=lambda: {
            "style": "APA",
            "format": "author-date",
        }
    )

    figures: Dict[str, Any] = field(
        default_factory=lambda: {
            "caption_position": "below",
            "font_size": 10,
            "font": "Times New Roman",
        }
    )

    tables: Dict[str, Any] = field(
        default_factory=lambda: {
            "caption_position": "above",
            "font_size": 10,
            "font": "Times New Roman",
            "borders": True,
        }
    )

    abbreviations: Dict[str, str] = field(default_factory=dict)

    @classmethod
    def from_dict(cls, data: Dict) -> "FormatConfig":
        config = cls()
        if "name" in data:
            config.name = data["name"]
        if "version" in data:
            config.version = data["version"]
        if "font" in data:
            config.font.update(data["font"])
        if "spacing" in data:
            config.spacing.update(data["spacing"])
        if "margins" in data:
            config.margins.update(data["margins"])
        if "references" in data:
            config.references.update(data["references"])
        if "figures" in data:
            config.figures.update(data["figures"])
        if "tables" in data:
            config.tables.update(data["tables"])
        if "abbreviations" in data:
            config.abbreviations = data["abbreviations"]
        return config


class FormatAdjuster:
    """formatter"""

    BUILTIN_FORMATS = {
        "nature": {
            "name": "Nature",
            "font": {
                "body": "Times New Roman",
                "body_size": 10,
                "title": "Arial",
                "title_size": 14,
                "caption": "Times New Roman",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "single",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "Nature",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
            },
        },
        "cell": {
            "name": "Cell",
            "font": {
                "body": "Arial",
                "body_size": 11,
                "title": "Arial",
                "title_size": 18,
                "caption": "Arial",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "Cell",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
            },
        },
        "nejm": {
            "name": "NEJM",
            "font": {
                "body": "Times New Roman",
                "body_size": 11,
                "title": "Times New Roman",
                "title_size": 16,
                "caption": "Times New Roman",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 0,
                "indent": 0,
            },
            "references": {
                "style": "NEJM",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
            },
        },
        "lancet": {
            "name": "The Lancet",
            "font": {
                "body": "Times New Roman",
                "body_size": 10,
                "title": "Arial",
                "title_size": 14,
                "caption": "Times New Roman",
                "caption_size": 8,
            },
            "spacing": {
                "line_space": "single",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "Lancet",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 8,
            },
            "tables": {
                "caption_position": "above",
                "caption_size": 8,
                "borders": True,
            },
        },
        "plos": {
            "name": "PLOS",
            "font": {
                "body": "Times New Roman",
                "body_size": 12,
                "title": "Arial",
                "title_size": 18,
                "caption": "Times New Roman",
                "caption_size": 10,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 6,
                "indent": 0.5,
            },
            "references": {
                "style": "PLOS",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 10,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 10,
                "borders": True,
                "title_bold": True,
            },
        },
        "bmc": {
            "name": "BMC",
            "font": {
                "body": "Times New Roman",
                "body_size": 12,
                "title": "Times New Roman",
                "title_size": 18,
                "caption": "Times New Roman",
                "caption_size": 10,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 6,
                "indent": 0.5,
            },
            "references": {
                "style": "BMC",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 10,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 10,
                "borders": True,
                "title_bold": True,
            },
        },
        "pnas": {
            "name": "PNAS",
            "font": {
                "body": "Times New Roman",
                "body_size": 11,
                "title": "Arial",
                "title_size": 14,
                "caption": "Arial",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "single",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "PNAS",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
            },
        },
        "embo": {
            "name": "EMBO",
            "font": {
                "body": "Arial",
                "body_size": 10,
                "title": "Arial",
                "title_size": 14,
                "caption": "Arial",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "single",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "EMBO",
                "format": "author-date",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
                "title_bold": True,
            },
        },
        "jbc": {
            "name": "JBC",
            "font": {
                "body": "Arial",
                "body_size": 11,
                "title": "Arial",
                "title_size": 14,
                "caption": "Arial",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "single",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "JBC",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
                "title_bold": True,
            },
        },
        "cell_reports": {
            "name": "Cell Reports",
            "font": {
                "body": "Arial",
                "body_size": 11,
                "title": "Arial",
                "title_size": 18,
                "caption": "Arial",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "Cell",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
                "title_bold": True,
            },
        },
        "scientific_reports": {
            "name": "Scientific Reports",
            "font": {
                "body": "Times New Roman",
                "body_size": 12,
                "title": "Times New Roman",
                "title_size": 18,
                "caption": "Times New Roman",
                "caption_size": 10,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 6,
                "indent": 0.5,
            },
            "references": {
                "style": "Scientific Reports",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 10,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 10,
                "borders": True,
                "title_bold": True,
            },
        },
        "jimmunol": {
            "name": "J Immunol",
            "font": {
                "body": "Arial",
                "body_size": 11,
                "title": "Arial",
                "title_size": 14,
                "caption": "Arial",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "single",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "J Immunol",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
                "title_bold": True,
            },
        },
        "molecular_cell": {
            "name": "Molecular Cell",
            "font": {
                "body": "Arial",
                "body_size": 11,
                "title": "Arial",
                "title_size": 18,
                "caption": "Arial",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "Molecular Cell",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
                "title_bold": True,
            },
        },
        "faseb": {
            "name": "FASEB Journal",
            "font": {
                "body": "Times New Roman",
                "body_size": 12,
                "title": "Arial",
                "title_size": 14,
                "caption": "Times New Roman",
                "caption_size": 10,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 6,
                "indent": 0.5,
            },
            "references": {
                "style": "FASEB",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 10,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 10,
                "borders": True,
                "title_bold": True,
            },
        },
        "science": {
            "name": "Science",
            "font": {
                "body": "Times New Roman",
                "body_size": 11,
                "title": "Arial",
                "title_size": 14,
                "caption": "Times New Roman",
                "caption_size": 9,
            },
            "spacing": {
                "line_space": "single",
                "paragraph_space": 6,
                "indent": 0,
            },
            "references": {
                "style": "Science",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 9,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 9,
                "borders": True,
            },
        },
        "ieee": {
            "name": "IEEE",
            "font": {
                "body": "Times New Roman",
                "body_size": 10,
                "title": "Times New Roman",
                "title_size": 24,
                "caption": "Times New Roman",
                "caption_size": 8,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 0,
                "indent": 0,
            },
            "references": {
                "style": "IEEE",
                "format": "numbered",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 8,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 8,
                "borders": True,
            },
        },
        "apa": {
            "name": "APA",
            "font": {
                "body": "Times New Roman",
                "body_size": 12,
                "title": "Times New Roman",
                "title_size": 12,
                "caption": "Times New Roman",
                "caption_size": 10,
            },
            "spacing": {
                "line_space": "double",
                "paragraph_space": 0,
                "indent": 0.5,
            },
            "references": {
                "style": "APA",
                "format": "author-date",
            },
            "figures": {
                "caption_position": "below",
                "font_size": 10,
            },
            "tables": {
                "caption_position": "above",
                "font_size": 10,
                "borders": False,
            },
        },
    }

    def __init__(self, config: FormatConfig = None):
        self.config = config or FormatConfig()

    def load_config(self, config_path: str) -> FormatConfig:
        """Load configuration file"""
        path = Path(config_path)

        if path.suffix.lower() in [".yaml", ".yml"]:
            import yaml

            data = yaml.safe_load(path.read_text(encoding="utf-8"))
        elif path.suffix.lower() == ".json":
            import json

            data = json.loads(path.read_text(encoding="utf-8"))
        else:
            raise ValueError(f"Unsupported configuration file format: {path.suffix}")

        self.config = FormatConfig.from_dict(data)
        return self.config

    def load_builtin_format(self, format_name: str) -> FormatConfig:
        """Load built-in formats"""
        fmt = format_name.lower().replace(" ", "-")

        if fmt not in self.BUILTIN_FORMATS:
            for key, value in self.BUILTIN_FORMATS.items():
                if value["name"].lower() == format_name.lower():
                    fmt = key
                    break
            else:
                raise ValueError(f"Built-in format not found: {format_name}")

        self.config = FormatConfig.from_dict(self.BUILTIN_FORMATS[fmt])
        return self.config

    def apply_format(self, text: str) -> str:
        """Apply formatting to text (Markdown format)"""
        lines = text.split("\n")
        result = []

        for line in lines:
            result.append(self._format_line(line))

        return "\n".join(result)

    def _format_line(self, line: str) -> str:
        """Format a single line"""
        line = self._format_headings(line)
        line = self._format_figures(line)
        line = self._format_tables(line)
        line = self._format_references(line)
        line = self._format_abbreviations(line)
        return line

    def _format_headings(self, line: str) -> str:
        """Format title"""
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            hashes, title = heading_match.groups()
            return f"{hashes} {title}"
        return line

    def _format_figures(self, line: str) -> str:
        """Format chart"""
        return line

    def _format_tables(self, line: str) -> str:
        """Format table"""
        return line

    def _format_references(self, line: str) -> str:
        """Formatting reference citations"""
        return line

    def _format_abbreviations(self, line: str) -> str:
        """Format abbreviation"""
        for abbr, full in self.config.abbreviations.items():
            line = re.sub(
                rf"\b{abbr}\b(?!\s*\([A-Z]\))",
                f"{abbr} ({full})",
                line,
            )
        return line

    def validate_format(self, text: str) -> Dict[str, List[str]]:
        """Verify that the format meets the requirements"""
        issues = []

        reference_issues = self._validate_references(text)
        if reference_issues:
            issues.extend(reference_issues)

        return {
            "issues": issues,
            "is_valid": len(issues) == 0,
        }

    def _validate_references(self, text: str) -> List[str]:
        """Verify reference format"""
        issues = []
        style = self.config.references.get("style", "APA")

        if style == "IEEE":
            if not re.search(r"\[\d+\]", text):
                issues.append("Numbered reference in IEEE format not found")

        return issues

    def apply_to_docx(self, docx_path: str, output_path: str = None) -> str:
        """Apply formatting to a Word document"""
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document(docx_path)

            body_font = self.config.font.get("body", "Times New Roman")
            title_font = self.config.font.get("title", "Arial")
            body_size = self.config.font.get("body_size", 10)
            title_size = self.config.font.get("title_size", 18)

            for paragraph in doc.paragraphs:
                if paragraph.style:
                    style_name = paragraph.style.name.lower()

                    if "heading" in style_name or "title" in style_name:
                        if paragraph.style.font.name:
                            paragraph.style.font.name = title_font
                        paragraph.style.font.size = Pt(title_size)
                    else:
                        if paragraph.style.font.name:
                            paragraph.style.font.name = body_font
                        paragraph.style.font.size = Pt(body_size)

            output_path = output_path or docx_path.replace(".docx", "_formatted.docx")
            doc.save(output_path)
            return output_path

        except ImportError:
            raise ImportError("Need to install python-docx library: pip install python-docx")

    def apply_to_tex(self, tex_path: str, output_path: str = None) -> str:
        """Apply formatting to a LaTeX document"""
        content = Path(tex_path).read_text(encoding="utf-8")

        output_path = output_path or tex_path.replace(".tex", "_formatted.tex")
        Path(output_path).write_text(content, encoding="utf-8")
        return output_path

    def generate_format_report(self) -> Dict[str, Any]:
        """Generate format configuration report"""
        return {
            "name": self.config.name,
            "font": self.config.font,
            "spacing": self.config.spacing,
            "margins": self.config.margins,
            "references": self.config.references,
            "figures": self.config.figures,
            "tables": self.config.tables,
        }
