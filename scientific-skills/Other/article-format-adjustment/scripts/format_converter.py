#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""format conversion module
Supports mutual conversion between Word, Markdown, and LaTeX formats"""

import subprocess
import tempfile
from pathlib import Path
from typing import Dict, Optional, Tuple


class FormatConverter:
    """format converter"""

    SUPPORTED_FORMATS = {
        "docx": {
            "name": "Word Document",
            "extensions": [".docx", ".doc"],
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
        "md": {
            "name": "Markdown",
            "extensions": [".md", ".markdown", ".txt"],
            "mime": "text/markdown",
        },
        "tex": {
            "name": "LaTeX",
            "extensions": [".tex", ".latex"],
            "mime": "application/x-latex",
        },
    }

    def __init__(self):
        self.temp_dir = None

    def get_format(self, file_path: str) -> str:
        """Detect file format"""
        ext = Path(file_path).suffix.lower()

        for fmt, info in self.SUPPORTED_FORMATS.items():
            if ext in info["extensions"]:
                return fmt

        raise ValueError(f"Unsupported file format: {ext}")

    def detect_format(self, content: str) -> str:
        """Detect format based on content"""
        if content.startswith("\\documentclass") or content.startswith(
            "\\documentclass"
        ):
            return "tex"
        if content.startswith("#") or content.startswith("---"):
            return "md"
        return "md"

    def to_markdown(self, input_path: str) -> str:
        """Convert to Markdown format"""
        fmt = self.get_format(input_path)

        if fmt == "md":
            return Path(input_path).read_text(encoding="utf-8")

        elif fmt == "docx":
            return self._docx_to_markdown(input_path)

        elif fmt == "tex":
            return self._tex_to_markdown(input_path)

        raise ValueError(f"Unable to convert format: {fmt}")

    def from_markdown(
        self, markdown_text: str, output_path: str, target_format: str = None
    ):
        """Convert from Markdown to target format"""
        if target_format is None:
            target_format = self.get_format(output_path)

        if target_format == "md":
            Path(output_path).write_text(markdown_text, encoding="utf-8")
            return output_path

        elif target_format == "docx":
            return self._markdown_to_docx(markdown_text, output_path)

        elif target_format == "tex":
            return self._markdown_to_tex(markdown_text, output_path)

        raise ValueError(f"Unable to convert to format: {target_format}")

    def _docx_to_markdown(self, docx_path: str) -> str:
        """Word to Markdown"""
        try:
            from docx import Document

            doc = Document(docx_path)
            markdown_lines = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    markdown_lines.append("")
                    continue

                heading_level = self._detect_heading_level(para)
                if heading_level:
                    markdown_lines.append(f"{'#' * heading_level} {text}")
                else:
                    markdown_lines.append(text)

            markdown_lines.extend(self._extract_tables_from_docx(doc))

            return "\n".join(markdown_lines)

        except ImportError:
            raise ImportError("Need to install python-docx library: pip install python-docx")

    def _tex_to_markdown(self, tex_path: str) -> str:
        """LaTeX to Markdown"""
        try:
            import panflute as pf

            doc = pf.parse_document(tex_path)
            md_text = pf.convert_text(
                doc, input_format="latex", output_format="markdown"
            )

            return md_text

        except ImportError:
            raise ImportError("Need to install panflute library: pip install panflute")

    def _markdown_to_docx(self, markdown_text: str, output_path: str) -> str:
        """Markdown to Word"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            doc = Document()

            lines = markdown_text.split("\n")
            in_table = False
            current_table = None

            for line in lines:
                line = line.strip()

                if not line:
                    continue

                if line.startswith("|"):
                    if not in_table:
                        current_table = doc.add_table(rows=1, cols=0)
                        current_table.style = "Table Grid"
                        in_table = True

                    cells = [c.strip() for c in line.strip("|").split("|")]
                    if current_table:
                        if len(current_table.columns) == 0:
                            for _ in cells:
                                current_table.add_column(Inches(1.5))
                        row = current_table.add_row()
                        for i, cell_text in enumerate(cells):
                            if i < len(row.cells):
                                row.cells[i].text = cell_text
                    continue

                in_table = False
                current_table = None

                if line.startswith("#"):
                    parts = line.split(None, 1)
                    level = len(parts[0])
                    text = parts[1] if len(parts) > 1 else ""

                    heading = doc.add_heading(text, level=level - 1)
                    heading.style.font.size = Pt(14 if level <= 2 else 12)
                else:
                    doc.add_paragraph(line)

            doc.save(output_path)
            return output_path

        except ImportError:
            raise ImportError("Need to install python-docx library: pip install python-docx")

    def _markdown_to_tex(self, markdown_text: str, output_path: str) -> str:
        """Markdown to LaTeX"""
        tex_content = """\\documentclass{article}
\\usepackage[utf8]{inputenc}
\\usepackage{amsmath}
\\usepackage{graphicx}
\\usepackage{hyperref}
\\usepackage{booktabs}

\\begin{document}

"""

        lines = markdown_text.split("\n")
        in_table = False
        table_rows = []

        for line in lines:
            line = line.strip()

            if not line:
                continue

            if line.startswith("|"):
                if not in_table:
                    in_table = True
                    table_rows = []
                cells = [c.strip() for c in line.strip("|").split("|")]
                table_rows.append(" & ".join(cells))
                continue

            if in_table:
                tex_content += (
                    "\\begin{tabular}{|"
                    + "l|" * len(table_rows[0].split(" & "))
                    + "}\n"
                )
                tex_content += "\\hline\n"
                tex_content += " \\\\\n\\hline\n".join(table_rows)
                tex_content += "\\hline\n\\end{tabular}\n\n"
                in_table = False
                table_rows = []

            if line.startswith("#"):
                parts = line.split(None, 1)
                level = len(parts[0])
                text = parts[1] if len(parts) > 1 else ""

                if level == 1:
                    tex_content += f"\\section{{{text}}}\n\n"
                elif level == 2:
                    tex_content += f"\\subsection{{{text}}}\n\n"
                elif level == 3:
                    tex_content += f"\\subsubsection{{{text}}}\n\n"
            else:
                tex_content += f"{line}\n\n"

        tex_content += "\\end{document}"
        Path(output_path).write_text(tex_content, encoding="utf-8")
        return output_path

    def _detect_heading_level(self, paragraph) -> int:
        """Check if a paragraph is a title"""
        style_name = paragraph.style.name.lower() if paragraph.style else ""

        if "heading 1" in style_name or style_name == "title":
            return 1
        elif "heading 2" in style_name:
            return 2
        elif "heading 3" in style_name:
            return 3
        elif "heading 4" in style_name:
            return 4

        return 0

    def _extract_tables_from_docx(self, doc) -> list:
        """Extract tables from Word documents"""
        table_lines = []

        for table in doc.tables:
            if not table.rows:
                continue

            header = []
            for cell in table.rows[0].cells:
                header.append(cell.text.strip())

            if header:
                separator = " | ".join(["---"] * len(header))
                header_line = "| " + " | ".join(header) + " |"
                table_lines.append("")
                table_lines.append(header_line)
                table_lines.append(separator)

                for row in table.rows[1:]:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_lines.append("| " + " | ".join(row_data) + " |")

        return table_lines

    def convert(self, input_path: str, output_path: str) -> Tuple[str, str]:
        """Universal conversion interface"""
        input_format = self.get_format(input_path)
        output_format = self.get_format(output_path)

        if input_format == output_format:
            return input_path, output_path

        if input_format == "docx" and output_format == "md":
            md_text = self._docx_to_markdown(input_path)
            Path(output_path).write_text(md_text, encoding="utf-8")
            return input_path, output_path

        if input_format == "md" and output_format == "docx":
            md_text = Path(input_path).read_text(encoding="utf-8")
            self._markdown_to_docx(md_text, output_path)
            return input_path, output_path

        if input_format == "docx" and output_format == "tex":
            md_text = self._docx_to_markdown(input_path)
            self._markdown_to_tex(md_text, output_path)
            return input_path, output_path

        if input_format == "tex" and output_format == "docx":
            md_text = self._tex_to_markdown(input_path)
            self._markdown_to_docx(md_text, output_path)
            return input_path, output_path

        if input_format == "md" and output_format == "tex":
            md_text = Path(input_path).read_text(encoding="utf-8")
            self._markdown_to_tex(md_text, output_path)
            return input_path, output_path

        if input_format == "tex" and output_format == "md":
            md_text = self._tex_to_markdown(input_path)
            Path(output_path).write_text(md_text, encoding="utf-8")
            return input_path, output_path

        raise ValueError(f"Unsupported conversion: {input_format} -> {output_format}")

    def get_pandoc_version(self) -> str:
        """Get pandoc version"""
        try:
            result = subprocess.run(
                ["pandoc", "--version"],
                capture_output=True,
                text=True,
                timeout=10,
            )
            return result.stdout.split("\n")[0]
        except Exception as e:
            return f"Pandoc not found: {e}"

    def __del__(self):
        if self.temp_dir and Path(self.temp_dir).exists():
            import shutil

            shutil.rmtree(self.temp_dir, ignore_errors=True)
