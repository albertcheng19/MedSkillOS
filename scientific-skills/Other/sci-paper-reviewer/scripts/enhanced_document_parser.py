#!/usr/bin/env python3
"""
SCI Paper Reviewer - Enhanced Document Parser

Enhanced version with actual PDF and Word document parsing capabilities.
This module provides production-ready document parsing for the SCI paper reviewer.
Optimized with robust pypdf extraction logic.
"""

import os
import sys
import logging
import json
from pathlib import Path
from typing import Dict, Optional, Union, List
import importlib.util

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class EnhancedDocumentParser:
    """Enhanced document parser with actual PDF and Word parsing capabilities."""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt'}
        self._check_dependencies()
        
    def _check_dependencies(self):
        """Check which parsing libraries are available."""
        self.has_pypdf = importlib.util.find_spec("pypdf") is not None
        self.has_pypdf2 = importlib.util.find_spec("PyPDF2") is not None
        self.has_pdfplumber = importlib.util.find_spec("pdfplumber") is not None
        self.has_pymupdf = importlib.util.find_spec("fitz") is not None
        self.has_docx = importlib.util.find_spec("docx") is not None
        
        logger.info(f"Available parsers - pypdf: {self.has_pypdf}, "
                   f"pdfplumber: {self.has_pdfplumber}, "
                   f"PyMuPDF: {self.has_pymupdf}, "
                   f"PyPDF2: {self.has_pypdf2}, "
                   f"python-docx: {self.has_docx}")
    
    def parse_document(self, file_path: Union[str, Path]) -> Dict[str, Union[str, bool]]:
        """
        Parse a document and extract text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing parsing results
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                'success': False,
                'content': '',
                'file_type': 'unknown',
                'error': f'File not found: {file_path}'
            }
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            return {
                'success': False,
                'content': '',
                'file_type': file_extension,
                'error': f'Unsupported file format: {file_extension}'
            }
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf_enhanced(file_path)
            elif file_extension == '.docx':
                return self._parse_docx_enhanced(file_path)
            elif file_extension == '.doc':
                return self._parse_doc_legacy(file_path)
            elif file_extension == '.txt':
                return self._parse_txt(file_path)
            else:
                return {
                    'success': False,
                    'content': '',
                    'file_type': file_extension,
                    'error': f'Parser not implemented for: {file_extension}'
                }
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            return {
                'success': False,
                'content': '',
                'file_type': file_extension,
                'error': f'Parsing error: {str(e)}'
            }
    
    def _parse_pdf_enhanced(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Parse PDF file using available libraries."""
        content = ""
        
        # Priority 1: pdfplumber (best for complex layouts)
        if self.has_pdfplumber:
            try:
                content = self._extract_with_pdfplumber(file_path)
                if content.strip():
                    logger.info(f"Successfully parsed PDF with pdfplumber")
                    return {
                        'success': True,
                        'content': content,
                        'file_type': 'pdf',
                        'error': '',
                        'parser_used': 'pdfplumber'
                    }
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
        
        # Priority 2: pypdf (modern, robust)
        if self.has_pypdf and not content.strip():
            try:
                content = self._extract_with_pypdf(file_path)
                if content.strip():
                    logger.info(f"Successfully parsed PDF with pypdf")
                    return {
                        'success': True,
                        'content': content,
                        'file_type': 'pdf',
                        'error': '',
                        'parser_used': 'pypdf'
                    }
            except Exception as e:
                logger.warning(f"pypdf failed: {e}")

        # Priority 3: PyMuPDF (fast, good backup)
        if self.has_pymupdf and not content.strip():
            try:
                content = self._extract_with_pymupdf(file_path)
                if content.strip():
                    logger.info(f"Successfully parsed PDF with PyMuPDF")
                    return {
                        'success': True,
                        'content': content,
                        'file_type': 'pdf',
                        'error': '',
                        'parser_used': 'pymupdf'
                    }
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")
        
        # Priority 4: PyPDF2 (legacy backup)
        if self.has_pypdf2 and not content.strip():
            try:
                content = self._extract_with_pypdf2(file_path)
                if content.strip():
                    logger.info(f"Successfully parsed PDF with PyPDF2")
                    return {
                        'success': True,
                        'content': content,
                        'file_type': 'pdf',
                        'error': '',
                        'parser_used': 'pypdf2'
                    }
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")
        
        # Fallback to basic implementation
        if not content.strip():
            content = self._extract_pdf_basic(file_path)
            return {
                'success': bool(content.strip()),
                'content': content,
                'file_type': 'pdf',
                'error': '' if content.strip() else 'No text content found in PDF',
                'parser_used': 'basic'
            }
        
        return {
            'success': False,
            'content': '',
            'file_type': 'pdf',
            'error': 'No PDF parsing libraries available'
        }
    
    def _parse_docx_enhanced(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Parse DOCX file using python-docx."""
        if self.has_docx:
            try:
                import docx
                doc = docx.Document(file_path)
                
                content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                content.append(cell.text.strip())
                
                full_content = '\n\n'.join(content)
                
                return {
                    'success': bool(full_content.strip()),
                    'content': full_content,
                    'file_type': 'docx',
                    'error': '',
                    'parser_used': 'python-docx'
                }
            except Exception as e:
                logger.error(f"python-docx parsing failed: {e}")
                return self._parse_docx_basic(file_path)
        else:
            return self._parse_docx_basic(file_path)
    
    def _parse_doc_legacy(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Parse legacy DOC files."""
        # Note: .doc files are binary format and require special handling
        # For now, provide basic implementation
        logger.warning("Legacy .doc format support is limited. Consider converting to .docx")
        return self._parse_docx_basic(file_path)
    
    def _parse_txt(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Parse plain text file."""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                'success': True,
                'content': content,
                'file_type': 'txt',
                'error': '',
                'parser_used': 'utf-8'
            }
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    
                    return {
                        'success': True,
                        'content': content,
                        'file_type': 'txt',
                        'error': '',
                        'parser_used': encoding,
                        'warning': f'Used {encoding} encoding instead of UTF-8'
                    }
                except UnicodeDecodeError:
                    continue
            
            return {
                'success': False,
                'content': '',
                'file_type': 'txt',
                'error': 'Unable to decode text file with any supported encoding'
            }
        except Exception as e:
            return {
                'success': False,
                'content': '',
                'file_type': 'txt',
                'error': f'Text file parsing error: {str(e)}'
            }
    
    def _extract_with_pdfplumber(self, file_path: Path) -> str:
        """Extract text using pdfplumber."""
        import pdfplumber
        content = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    content.append(f"=== Page {page_num} ===\n{page_text}")
        
        return '\n\n'.join(content)

    def _extract_with_pypdf(self, file_path: Path) -> str:
        """Extract text using pypdf."""
        import pypdf
        content = []
        reader = pypdf.PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                content.append(f"=== Page {i+1} ===\n{text}\n")
        return "\n".join(content)
    
    def _extract_with_pymupdf(self, file_path: Path) -> str:
        """Extract text using PyMuPDF."""
        import fitz  # PyMuPDF
        content = []
        
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc, 1):
            page_text = page.get_text()
            if page_text.strip():
                content.append(f"=== Page {page_num} ===\n{page_text.strip()}")
        
        doc.close()
        return '\n\n'.join(content)
    
    def _extract_with_pypdf2(self, file_path: Path) -> str:
        """Extract text using PyPDF2."""
        import PyPDF2
        content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    content.append(f"=== Page {page_num} ===\n{page_text}")
        
        return '\n\n'.join(content)
    
    def _extract_pdf_basic(self, file_path: Path) -> str:
        """
        Basic PDF extraction (fallback) using only standard libraries.
        Attempts to extract text streams and decompress them using zlib.
        This is a best-effort approach and will likely be imperfect.
        """
        logger.warning("Using basic PDF extraction (Standard Library Only). Quality may be low.")
        
        import zlib
        import re
        
        text_content = []
        
        try:
            with open(file_path, 'rb') as f:
                data = f.read()
                
            # Pattern to find streams
            # PDF streams are usually between 'stream' and 'endstream' markers
            stream_pattern = re.compile(b'stream(.*?)endstream', re.DOTALL)
            
            for match in stream_pattern.finditer(data):
                stream_data = match.group(1).strip()
                
                # Try to decompress
                try:
                    # PDF streams are often compressed with zlib (FlateDecode)
                    decompressed = zlib.decompress(stream_data)
                    
                    # Extract text from decompressed stream
                    # Text is usually in parentheses like (Hello World)Tj or [ (Hello) (World) ]TJ
                    # This regex looks for text inside parentheses
                    text_matches = re.findall(b'\((.*?)\)', decompressed)
                    
                    for text in text_matches:
                        try:
                            # Try to decode as UTF-8, fallback to Latin-1
                            decoded_text = text.decode('utf-8', errors='ignore')
                            # Simple filter to avoid extracting garbage
                            if len(decoded_text) > 1 and not decoded_text.startswith('\\'):
                                # Handle escaped characters simply
                                decoded_text = decoded_text.replace('\\n', '\n').replace('\\(', '(').replace('\\)', ')')
                                text_content.append(decoded_text)
                        except:
                            pass
                            
                except zlib.error:
                    # Not a zlib compressed stream or corrupted
                    continue
                    
            if not text_content:
                return f"[Basic Parser] Could not extract readable text from {file_path.name}. The PDF might be encrypted or use unsupported encoding. Please install pypdf, pdfplumber, or PyMuPDF."
                
            return "[Basic Parser Result - Formatting Lost]\n\n" + " ".join(text_content)
            
        except Exception as e:
            return f"[Basic Parser Error] Failed to parse {file_path.name}: {str(e)}"
    
    def _parse_docx_basic(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Basic DOCX extraction (fallback) using standard libraries."""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            content = []
            
            with zipfile.ZipFile(file_path) as docx:
                if 'word/document.xml' in docx.namelist():
                    xml_content = docx.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    
                    # Define namespace
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    # Find all text nodes
                    for p in root.findall('.//w:p', ns):
                        texts = [node.text for node in p.findall('.//w:t', ns) if node.text]
                        if texts:
                            content.append(''.join(texts))
            
            full_content = '\n\n'.join(content)
            
            if not full_content:
                full_content = f"[Basic Parser] Could not extract text from {file_path.name}."
            
            return {
                'success': True,
                'content': full_content,
                'file_type': 'docx',
                'error': '',
                'parser_used': 'basic_xml',
                'warning': 'Using basic XML parsing. Formatting may be lost.'
            }
            
        except Exception as e:
            return {
                'success': False,
                'content': '',
                'file_type': 'docx',
                'error': f'Basic DOCX parsing error: {str(e)}'
            }
    
    def extract_sections(self, content: str) -> Dict[str, str]:
        """
        Extract key sections from academic paper content.
        
        Args:
            content: Full text content of the document
            
        Returns:
            Dict containing extracted sections
        """
        sections = {
            'abstract': '',
            'introduction': '',
            'methods': '',
            'results': '',
            'discussion': '',
            'conclusion': '',
            'references': '',
            'full_text': content
        }
        
        content_lower = content.lower()
        
        # Define section patterns and their variations
        section_patterns = {
            'abstract': ['abstract', 'summary'],
            'introduction': ['introduction', 'background'],
            'methods': ['methods', 'methodology', 'materials and methods', 'experimental procedures'],
            'results': ['results', 'findings'],
            'discussion': ['discussion'],
            'conclusion': ['conclusion', 'conclusions', 'summary and conclusion'],
            'references': ['references', 'bibliography', 'literature cited']
        }
        
        # Extract sections
        for section_name, patterns in section_patterns.items():
            for pattern in patterns:
                if pattern in content_lower:
                    try:
                        section_content = self._extract_section_content(content, content_lower, pattern)
                        if section_content and len(section_content.strip()) > 50:  # Minimum content threshold
                            sections[section_name] = section_content.strip()
                            break
                    except Exception as e:
                        logger.warning(f"Could not extract {section_name} section: {e}")
        
        return sections
    
    def _extract_section_content(self, original_content: str, lower_content: str, section_name: str) -> str:
        """Extract content for a specific section."""
        section_start = lower_content.find(section_name)
        if section_start == -1:
            return ""
        
        # Find the end of the section (next major section or end of document)
        possible_end_sections = [
            'abstract', 'introduction', 'methods', 'results', 
            'discussion', 'conclusion', 'references', 'acknowledgments'
        ]
        
        # Remove the current section from possible end sections
        possible_end_sections = [s for s in possible_end_sections if s != section_name]
        
        section_end = len(original_content)
        start_search_pos = section_start + len(section_name) + 10  # Skip past the section header
        
        for end_section in possible_end_sections:
            end_pos = lower_content.find(end_section, start_search_pos)
            if end_pos != -1 and end_pos < section_end:
                section_end = end_pos
        
        # Extract the content
        section_content = original_content[section_start:section_end].strip()
        
        # Remove the section header line
        lines = section_content.split('\n')
        if lines and section_name.lower() in lines[0].lower():
            section_content = '\n'.join(lines[1:]).strip()
        
        return section_content
    
    def get_document_info(self, file_path: Union[str, Path]) -> Dict[str, Union[str, int, bool]]:
        """
        Get comprehensive information about the document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing document information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                'exists': False,
                'file_name': file_path.name,
                'file_size': 0,
                'file_type': 'unknown'
            }
        
        file_stats = file_path.stat()
        
        info = {
            'exists': True,
            'file_name': file_path.name,
            'file_size': file_stats.st_size,
            'file_type': file_path.suffix.lower(),
            'modified_time': file_stats.st_mtime,
            'created_time': file_stats.st_ctime,
            'human_readable_size': self._format_file_size(file_stats.st_size)
        }
        
        # Add parsing capability info
        file_extension = file_path.suffix.lower()
        if file_extension == '.pdf':
            info['parsing_capabilities'] = {
                'pdfplumber': self.has_pdfplumber,
                'pypdf': self.has_pypdf,
                'pymupdf': self.has_pymupdf,
                'pypdf2': self.has_pypdf2
            }
        elif file_extension == '.docx':
            info['parsing_capabilities'] = {
                'python-docx': self.has_docx
            }
        
        return info
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"


def main():
    """Enhanced main function with better error handling."""
    if len(sys.argv) < 2:
        print("Usage: python enhanced_document_parser.py <file_path> [output_file_or_format]")
        print("Arguments:")
        print("  file_path: Path to the document to parse")
        print("  output_file_or_format: Optional. Can be 'json', 'text', 'sections' OR a file path (e.g., .txt)")
        sys.exit(1)
    
    file_path = sys.argv[1]
    second_arg = sys.argv[2] if len(sys.argv) > 2 else 'json'
    
    # Check if second argument is a file path for output
    output_to_file = False
    output_file_path = None
    output_format = 'json'
    
    if second_arg not in ['json', 'text', 'sections']:
        # Assume it's a file path
        output_to_file = True
        output_file_path = second_arg
        output_format = 'text' # Force text format when writing to file directly
    else:
        output_format = second_arg
    
    parser = EnhancedDocumentParser()
    
    # Get document info first
    doc_info = parser.get_document_info(file_path)
    
    if not doc_info['exists']:
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
    
    # Parse document
    result = parser.parse_document(file_path)
    
    if result['success']:
        # If explicitly outputting to file (new simple mode)
        if output_to_file and output_file_path:
            try:
                with open(output_file_path, 'w', encoding='utf-8') as f:
                    f.write(result['content'])
                print(f"Successfully extracted text to {output_file_path}")
                # Log warning if empty
                if not result['content'].strip():
                     print("Warning: No text extracted. The file might be empty or scanned.", file=sys.stderr)
            except Exception as e:
                print(f"Error writing to output file: {e}", file=sys.stderr)
                sys.exit(1)
            sys.exit(0)

        # Legacy/Structured output mode
        print(f"Document: {doc_info['file_name']} ({doc_info.get('human_readable_size', 'Unknown size')})")
        print(f"✓ Successfully parsed {result['file_type']} document using {result.get('parser_used', 'unknown parser')}")
        print(f"✓ Extracted {len(result['content'])} characters of text")
        
        # Extract sections
        sections = parser.extract_sections(result['content'])
        
        # Output based on format
        if output_format == 'text':
            print(f"\n{'='*50}")
            print("FULL EXTRACTED CONTENT:")
            print(f"{'='*50}")
            print(result['content'])
            
        elif output_format == 'sections':
            print(f"\n{'='*50}")
            print("EXTRACTED SECTIONS:")
            print(f"{'='*50}")
            for section, content in sections.items():
                if section != 'full_text' and content:
                    print(f"\n--- {section.upper()} ---")
                    print(content[:500] + "..." if len(content) > 500 else content)
            
        else:  # json (default)
            output_data = {
                'document_info': doc_info,
                'parsing_result': result,
                'sections': sections,
                'extraction_timestamp': logger.handlers[0].baseFilename if hasattr(logger.handlers[0], 'baseFilename') else 'unknown'
            }
            
            output_file = Path(file_path).with_suffix('.parsed.json')
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(output_data, f, indent=2, ensure_ascii=False)
            
            print(f"\n✓ Parsed data saved to: {output_file}")
            
            # Also create a text summary
            summary_file = Path(file_path).with_suffix('.summary.txt')
            with open(summary_file, 'w', encoding='utf-8') as f:
                f.write(f"Document: {doc_info['file_name']}\n")
                f.write(f"Type: {result['file_type']}\n")
                f.write(f"Parser: {result.get('parser_used', 'unknown')}\n")
                f.write(f"Content length: {len(result['content'])} characters\n\n")
                
                f.write("EXTRACTED SECTIONS:\n")
                f.write("="*30 + "\n")
                for section, content in sections.items():
                    if section != 'full_text' and content:
                        f.write(f"\n{section.upper()}:\n")
                        f.write("-"*20 + "\n")
                        f.write(content[:1000] + "...\n" if len(content) > 1000 else content + "\n")
            
            print(f"✓ Summary saved to: {summary_file}")
    
    else:
        print(f"✗ Error parsing document: {result['error']}")
        sys.exit(1)


if __name__ == "__main__":
    main()
