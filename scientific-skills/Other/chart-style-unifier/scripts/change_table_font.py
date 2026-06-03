"""Table/Graph Font Unification Script

Usage:
1. python scripts/init_run.py --table (initialization)
2. AI fills in outputs/runs/<timestamp>/config.json
3. python scripts/change_table_font.py (execution)

Path constraints:
- Only allow access to user-supplied Word document paths
- The output directory is fixed to outputs/runs/<timestamp>/"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

ALLOWED_OUTPUT_DIR = Path("outputs/runs")

try:
    from docx import Document
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    qn = None
    DOCX_AVAILABLE = False

try:
    import pythoncom
    import win32com.client as win32

    WIN32_AVAILABLE = True
except ImportError:
    pythoncom = None
    win32 = None
    WIN32_AVAILABLE = False


def find_latest_run_dir():
    """Find the latest running directory"""
    runs_dir = ALLOWED_OUTPUT_DIR
    if not runs_dir.exists():
        return None

    run_dirs = sorted(runs_dir.iterdir(), reverse=True)
    for d in run_dirs:
        config_path = d / "config.json"
        if config_path.exists():
            return d
    return None


def load_config(run_dir: Path):
    """Load configuration file"""
    config_path = run_dir / "config.json"
    if not config_path.exists():
        print(f"[mistake] Configuration file does not exist")
        sys.exit(1)

    with config_path.open(encoding="utf-8") as f:
        return json.load(f)


def validate_doc_path(doc_path: str) -> bool:
    """Verify that the document path is safe"""
    if not doc_path:
        return False

    path = Path(doc_path)

    if not path.exists():
        return False

    if not str(path).lower().endswith(".docx"):
        return False

    return True


def change_table_font(config: dict):
    """Modify Word table font"""
    if not DOCX_AVAILABLE:
        print("[Tip] The python-docx library is missing, trying to install it...")
        import subprocess

        result = subprocess.run(
            ["pip", "install", "python-docx"], capture_output=True, text=True
        )
        if result.returncode != 0:
            print("[Error] Failed to install python-docx, please execute manually: pip install python-docx")
            sys.exit(1)

    doc_path = config["target_doc_path"]
    font_name = config["font_name"]
    font_size = config.get("font_size")

    if not doc_path:
        print("[Error] Please set target_doc_path in config.json as the target document path")
        sys.exit(1)

    if not validate_doc_path(doc_path):
        print(f"[mistake] The document path is invalid or the file does not exist: {doc_path}")
        sys.exit(1)

    if not font_name:
        print("[Error] Please set font_name as the target font name in config.json")
        sys.exit(1)

    try:
        doc = Document(doc_path)
    except Exception:
        print(f"[mistake] Unable to open document，Please confirm that the file is not occupied by other programs: {doc_path}")
        sys.exit(1)

    tables = doc.tables

    if not tables:
        print("[Tip] There is no table in the document, no need to process it")
        return

    print(f"Modifying {len(tables)} The font of the table is {font_name}...")

    table_count = 0
    cell_count = 0

    for table in tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                        if font_size is not None:
                            run.font.size = font_size

    # save
    doc.save(doc_path)

    print("[Finish]")
    print(f"- Modified {table_count} tables")
    print(f"- Modified approx. {cell_count} cells")
    print(f"- font: {font_name}")
    if font_size is not None:
        print(f"- Font size: {font_size}")


def set_table_numbers_italic(config: dict):
    """Italicize numbers in table"""
    if not DOCX_AVAILABLE:
        print("[Tip] The python-docx library is missing, trying to install it...")
        import subprocess

        result = subprocess.run(
            ["pip", "install", "python-docx"], capture_output=True, text=True
        )
        if result.returncode != 0:
            print("[Error] Failed to install python-docx, please execute manually: pip install python-docx")
            sys.exit(1)

    doc_path = config["target_doc_path"]

    if not doc_path:
        print("[Error] Please set target_doc_path in config.json as the target document path")
        sys.exit(1)

    if not validate_doc_path(doc_path):
        print(f"[mistake] The document path is invalid or the file does not exist: {doc_path}")
        sys.exit(1)

    try:
        doc = Document(doc_path)
    except Exception:
        print(f"[mistake] Unable to open document，Please confirm that the file is not occupied by other programs: {doc_path}")
        sys.exit(1)

    tables = doc.tables

    if not tables:
        print("[Tip] There is no table in the document, no need to process it")
        return

    print("Italicizing numbers in table...")

    table_count = 0
    cell_count = 0
    paragraph_count = 0

    for table in tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if re.search(r"[\d]+\.?[\d]*", text):
                        for run in paragraph.runs:
                            run.font.italic = True
                        paragraph_count += 1

    # save
    doc.save(doc_path)

    print("[Finish]")
    print(f"- Modified {table_count} tables")
    print(f"- Modified approx. {cell_count} cells")
    print(f"- {paragraph_count} Paragraphs containing numbers at have been italicized")


def set_table_chinese_italic(config: dict):
    """Set Chinese characters in the table to italics"""
    if not DOCX_AVAILABLE:
        print("The python-docx library is missing, please install it first:")
        print("pip install python-docx")
        sys.exit(1)

    doc_path = config["target_doc_path"]

    if not doc_path:
        print("Error: target_doc_path in config.json is required")
        sys.exit(1)

    if not Path(doc_path).exists():
        print(f"mistake: Document does not exist - {doc_path}")
        sys.exit(1)

    doc = Document(doc_path)
    tables = doc.tables

    if not tables:
        print("There are no tables in the document")
        return

    print(f"Setting the Chinese characters in the table to italics...")

    table_count = 0
    cell_count = 0
    paragraph_count = 0

    for table in tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if re.search(r"[\u4e00-\u9fff]", text):
                        for run in paragraph.runs:
                            run.font.italic = True
                        paragraph_count += 1

    # save
    doc.save(doc_path)

    print("[Finish]")
    print(f"- Modified {table_count} tables")
    print(f"- Modified approx. {cell_count} cells")
    print(f"- {paragraph_count} Paragraphs containing Chinese characters have been italicized")


def set_table_letters_lowercase(config: dict):
    """Change the English letters in the table to lowercase"""
    if not DOCX_AVAILABLE:
        print("[Tip] The python-docx library is missing, trying to install it...")
        import subprocess

        result = subprocess.run(
            ["pip", "install", "python-docx"], capture_output=True, text=True
        )
        if result.returncode != 0:
            print("[Error] Failed to install python-docx, please execute manually: pip install python-docx")
            sys.exit(1)

    doc_path = config["target_doc_path"]

    if not doc_path:
        print("[Error] Please set target_doc_path in config.json as the target document path")
        sys.exit(1)

    if not validate_doc_path(doc_path):
        print(f"[mistake] The document path is invalid or the file does not exist: {doc_path}")
        sys.exit(1)

    try:
        doc = Document(doc_path)
    except Exception:
        print(f"[mistake] Unable to open document，Please confirm that the file is not occupied by other programs: {doc_path}")
        sys.exit(1)

    tables = doc.tables

    if not tables:
        print("[Tip] There is no table in the document, no need to process it")
        return

    print("Changing English letters in the table to lowercase...")

    table_count = 0
    cell_count = 0
    run_count = 0

    for table in tables:
        table_count += 1
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if re.search(r"[A-Z]", run.text):
                            run.text = run.text.lower()
                            run_count += 1

    # save
    doc.save(doc_path)

    print("[Finish]")
    print(f"- Modified {table_count} tables")
    print(f"- Modified approx. {cell_count} cells")
    print(f"- {run_count} English letters have been changed to lowercase")


def apply_font(target_font, font_name: str, font_size):
    """Set font name and size as much as possible"""
    if target_font is None:
        return

    try:
        target_font.Name = font_name
    except Exception:
        pass

    try:
        target_font.NameFarEast = font_name
    except Exception:
        pass

    if font_size is not None:
        try:
            target_font.Size = font_size
        except Exception:
            pass


def apply_chart_font(chart, font_name: str, font_size):
    """Apply chart fonts to common elements"""
    try:
        apply_font(
            chart.ChartArea.Format.TextFrame2.TextRange.Font, font_name, font_size
        )
    except Exception:
        pass

    try:
        if chart.HasTitle:
            apply_font(
                chart.ChartTitle.Format.TextFrame2.TextRange.Font, font_name, font_size
            )
    except Exception:
        pass

    try:
        if chart.HasLegend:
            apply_font(
                chart.Legend.Format.TextFrame2.TextRange.Font, font_name, font_size
            )
    except Exception:
        pass

    for axis_type in (1, 2):
        for axis_group in (1, 2):
            axis = None
            try:
                axis = chart.Axes(axis_type, axis_group)
            except Exception:
                axis = None

            if axis is None:
                continue

            try:
                apply_font(axis.TickLabels.Font, font_name, font_size)
            except Exception:
                pass

            try:
                if axis.HasTitle:
                    apply_font(
                        axis.AxisTitle.Format.TextFrame2.TextRange.Font,
                        font_name,
                        font_size,
                    )
            except Exception:
                pass

    try:
        series_collection = chart.SeriesCollection()
        series_count = series_collection.Count
        for i in range(1, series_count + 1):
            series = series_collection(i)
            try:
                if series.HasDataLabels:
                    apply_font(series.DataLabels.Font, font_name, font_size)
            except Exception:
                pass
    except Exception:
        pass


def change_chart_font(config: dict):
    """Modify Word chart font"""
    if not WIN32_AVAILABLE:
        print("[Tip] The pywin32 library is missing, trying to install it...")
        import subprocess

        result = subprocess.run(
            ["pip", "install", "pywin32"], capture_output=True, text=True
        )
        if result.returncode != 0:
            print("[Error] Failed to install pywin32, please execute manually: pip install pywin32")
            sys.exit(1)

    doc_path = config["target_doc_path"]
    font_name = config["font_name"]
    font_size = config.get("font_size")

    if not doc_path:
        print("[Error] Please set target_doc_path in config.json as the target document path")
        sys.exit(1)

    if not validate_doc_path(doc_path):
        print(f"[mistake] The document path is invalid or the file does not exist: {doc_path}")
        sys.exit(1)

    if not font_name:
        print("[Error] Please set font_name as the target font name in config.json")
        sys.exit(1)

    pythoncom.CoInitialize()
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    doc = None
    chart_count = 0

    try:
        doc = word.Documents.Open(str(Path(doc_path).resolve()))

        for i in range(1, doc.InlineShapes.Count + 1):
            inline_shape = doc.InlineShapes(i)
            if inline_shape.HasChart:
                apply_chart_font(inline_shape.Chart, font_name, font_size)
                chart_count += 1

        for i in range(1, doc.Shapes.Count + 1):
            shape = doc.Shapes(i)
            if shape.HasChart:
                apply_chart_font(shape.Chart, font_name, font_size)
                chart_count += 1

        doc.Save()
    except Exception as e:
        print(f"[mistake] An error occurred while processing the chart: {str(e)}")
    finally:
        if doc is not None:
            doc.Close(SaveChanges=False)
        word.Quit()
        pythoncom.CoUninitialize()

    if chart_count > 0:
        print("[Finish]")
        print(f"- Modified {chart_count} charts")
        print(f"- font: {font_name}")
        if font_size is not None:
            print(f"- Font size: {font_size}")
    else:
        print("[Tip] There are no editable charts in the document, no need to process")


def main():
    """main entrance"""
    # Find the latest running directory
    run_dir = find_latest_run_dir()
    if not run_dir:
        print("[Error] Run directory not found, please run first: python scripts/init_run.py --table")
        sys.exit(1)

    print(f"Use configuration directory: {run_dir}")

    # Load configuration
    config = load_config(run_dir)

    # check status
    if config.get("status") == "done":
        print("[Tip] This configuration has been executed before. If you need to execute it again, please change the status to pending.")
        sys.exit(0)

    # implement
    target_type = config.get("target_type", "table")
    italic_numbers = config.get("italic_numbers", False)
    italic_chinese = config.get("italic_chinese", False)
    lowercase_letters = config.get("lowercase_letters", False)

    if lowercase_letters:
        set_table_letters_lowercase(config)
    elif italic_chinese:
        set_table_chinese_italic(config)
    elif italic_numbers:
        set_table_numbers_italic(config)
    elif target_type == "table":
        change_table_font(config)
    elif target_type == "chart":
        change_chart_font(config)
    else:
        print(f"[mistake] Not supported yet target_type={target_type}，Currently only supports table/chart")
        sys.exit(1)

    # update status
    config["status"] = "done"
    with (run_dir / "config.json").open("w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()
