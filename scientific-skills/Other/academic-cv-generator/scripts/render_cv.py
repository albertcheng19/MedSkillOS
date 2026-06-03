#!/usr/bin/env python
import argparse
import random
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


CN_HEADERS = [
    "personal data",
    "Basic information",
    "Educational background",
    "educate",
    "Work experience",
    "Project experience",
    "Participate in projects",
    "Paper published",
    "Publish a paper",
    "Published literature",
    "Awards",
    "Award-winning",
    "Skill",
    "master skills",
]
EN_HEADERS = [
    "Summary",
    "Education",
    "Work Experience",
    "Projects",
    "Publications",
    "Awards",
    "Skills",
    "Interests",
]

HEADER_MAP = {
    "profile": ("personal data", "Profile"),
    "education": ("Educational background", "Education"),
    "work": ("Work experience", "Work Experience"),
    "projects": ("Project experience", "Projects"),
    "publications": ("Paper published", "Publications"),
    "awards": ("Awards", "Awards"),
    "skills": ("master skills", "Skills"),
}

DEGREE_KEYWORDS = ("PhD", "master", "Bachelor", "PhD", "Master", "Bachelor")
PROJECT_KEYWORDS = ("Project name", "project", "platform", "system", "develop")
AWARD_KEYWORDS = ("Award-winning", "award", "funding", "fund", "best paper")
SKILL_KEYWORDS = ("Skill", "Skills", "Skill", "programming language", "Technical field", "Tools & Techniques", "tool")
WORK_KEYWORDS = ("Work", "hold office", "Taking office", "take charge of", "company")
EDU_KEYWORDS = ("educate", "Bachelor of Science")
PUB_HINTS = ("IEEE", "ACM", "Journal", "Proceedings", "Transactions")

COLOR_MAP = {
    "purple": RGBColor(0x6A, 0x1B, 0x9A),
    "cyan": RGBColor(0x00, 0xBC, 0xD4),
    "green": RGBColor(0x2E, 0x7D, 0x32),
    "red": RGBColor(0xC6, 0x28, 0x28),
}

DEGREE_RANK = {"PhD": 3, "master": 2, "Bachelor": 1}


def detect_lang(text):
    return "zh" if re.search(r"[\u4e00-\u9fff]", text) else "en"


def clean_lines(text):
    lines = [ln.lstrip("\ufeff").strip() for ln in text.splitlines()]
    return [ln for ln in lines if ln]


def strip_leading_markers(line):
    line = re.sub(r"^\s*[\-•*]\s*", "", line)
    line = re.sub(r"^\s*\d+[.)]\s*", "", line)
    line = line.strip("*").strip()
    if line.startswith("**") and line.endswith("**"):
        line = line.strip("*").strip()
    return line.strip()


def normalize_header_key(line):
    key = strip_leading_markers(line)
    key = key.strip().rstrip(":：").strip()
    return key


def has_explicit_headers(lines):
    for ln in lines:
        key = normalize_header_key(ln)
        if key in CN_HEADERS or key in EN_HEADERS:
            return True
        if key in ("Basic information", "Educational background", "Project experience", "Paper published", "Publish a paper", "Awards", "Skill", "master skills") and len(key) <= 6:
            return True
    return False


def split_sentences(text):
    text = text.replace("\n", " ")
    def protect_email(m):
        return m.group(0).replace(".", "<DOT>")
    text = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", protect_email, text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    parts = re.split(r"(?<=[。！？.!?])\s*", text)
    parts = [p.replace("<DOT>", ".") for p in parts]
    return [p.strip() for p in parts if p.strip()]


def is_contact_line(line):
    if re.fullmatch(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", line):
        return True
    if re.fullmatch(r"(\+?\d[\d\- ]{7,}\d)", line):
        return True
    if re.match("^(Name|Position|Title|School|Unit|Institution|Institution|E-mail|E-mail|Telephone)[::]", line):
        return True
    return False


def is_publication_line(line):
    if re.search(r"\(\d{4}\)", line):
        return True
    if any(k in line for k in PUB_HINTS):
        return True
    if re.search(r"\b(19|20)\d{2}\b", line) and ("," in line or "《" in line):
        return True
    return False


def is_award_line(line):
    if any(k in line for k in AWARD_KEYWORDS) and re.search(r"\b(19|20)\d{2}\b", line):
        return True
    return False


def is_education_line(line):
    if any(k in line for k in DEGREE_KEYWORDS):
        return True
    if any(k in line for k in EDU_KEYWORDS):
        return True
    return False


def is_work_line(line):
    return any(k in line for k in WORK_KEYWORDS)


def is_skill_line(line):
    if any(line.startswith(k) for k in SKILL_KEYWORDS):
        return True
    return any(k in line for k in ("skilled", "master", "programming language", "machine learning", "deep learning", "data mining"))


def polish_line(line):
    line = re.sub(r"\s{2,}", " ", line).strip()
    line = line.replace(" ：", "：").replace("： ", "：")
    line = line.replace(" ,", ",").replace(" .", ".")
    return line


def score_sections(line, current):
    scores = {k: 0 for k in HEADER_MAP.keys()}

    if is_publication_line(line) or "《" in line or "paper" in line:
        scores["publications"] += 4
    if is_award_line(line) or any(k in line for k in ("Award-winning", "award", "funding", "fund")):
        scores["awards"] += 4
    if is_education_line(line):
        scores["education"] += 3
    if any(k in line for k in PROJECT_KEYWORDS) or re.search("\\d{4}year", line):
        scores["projects"] += 3
    if "time" in line or "describe" in line or "Project name" in line:
        scores["projects"] += 2
    if is_skill_line(line):
        scores["skills"] += 3
    if is_work_line(line):
        scores["work"] += 3

    if current and current in scores:
        scores[current] += 1

    return scores


def classify_section(line, current, summary_set):
    header_key = normalize_header_key(line)
    if header_key in CN_HEADERS + EN_HEADERS:
        if "personal" in header_key or "Basic information" in header_key or "Profile" in header_key:
            return "profile"
        if "educate" in header_key or "Education" in header_key:
            return "education"
        if "Work" in header_key or "Work" in header_key:
            return "work"
        if "project" in header_key or "Project" in header_key:
            return "projects"
        if "paper" in header_key or "literature" in header_key or "Publication" in header_key:
            return "publications"
        if "Award-winning" in header_key or "Award" in header_key:
            return "awards"
        if "Skill" in header_key or "Skills" in header_key:
            return "skills"

    scores = score_sections(line, current)
    best = max(scores, key=lambda k: scores[k])
    if scores[best] >= 3:
        return best
    if current:
        return current
    if not summary_set:
        return "profile"
    return ""


def dedupe_keep_order(items):
    seen = set()
    out = []
    for it in items:
        if it in seen:
            continue
        seen.add(it)
        out.append(it)
    return out


def extract_org_from_text(text):
    m = re.search("(?:studied in|at|at)?([\\u4e00-\\u9fff]{2,30}(University|College|Institute|Department))", text)
    if not m:
        return ""
    org = m.group(1)
    org = re.sub("^(obtained in the year|obtained|currently studied in|at|at)", "", org)
    return org


def extract_major(sentence):
    # Prefer explicit majors; avoid matching degree words like "doctoral degree"
    patterns = [
        "(Computer Science|Software Engineering|Biomedicine|Biological Sciences|Data Science|Artificial Intelligence|Life Sciences|Information Technology)",
        "([\\u4e00-\\u9fff]{2,10})(?:Professional|Direction|Field)",
        "([\\u4e00-\\u9fff]{2,10}(?:Engineering|Science|Technology|Medicine|Biology))",
    ]
    for pat in patterns:
        m = re.search(pat, sentence)
        if m:
            return m.group(1)
    return ""


def extract_narrative_education(sentences, org):
    items = []
    default_major = ""
    for sent in sentences:
        if not default_major and ("field" in sent or "direction" in sent):
            default_major = extract_major(sent) or default_major
    for sent in sentences:
        year_matches = list(re.finditer("((?:19|20)\\d{2})year", sent))
        for i, ym in enumerate(year_matches):
            year = f"{ym.group(1)}Year"
            start = ym.start()
            end = year_matches[i + 1].start() if i + 1 < len(year_matches) else len(sent)
            segment = sent[start:end]
            degree_m = re.search("(Ph.D. | Master | Bachelor's) degree", segment)
            if not degree_m:
                continue
            degree = degree_m.group(1)
            school = extract_org_from_text(segment) or org
            major = ""
            major_m = re.search("Obtained the (?:(?P<school>[\\u4e00-\\u9fff]{2,30} University))?(?P<major>[\\u4e00-\\u9fff]{2,10})?(Doctoral | Master | Bachelor's) degree", segment)
            if major_m and major_m.group("major"):
                cand = major_m.group("major")
                if not re.search("(University|College|Institute|Department)", cand):
                    major = cand
            if not major:
                major = extract_major(segment) or default_major
            major = major.rstrip("of")
            parts = [degree]
            if major:
                parts.append(major)
            if school:
                parts.append(school)
            parts.append(year)
            items.append("，".join(parts))
    return dedupe_keep_order(items)


def extract_narrative_publications(sentences):
    items = []
    for sent in sentences:
        if not any(k in sent for k in ("paper", "publish", "《", '"')):
            continue
        cn_titles = re.findall(r"《([^》]+)》", sent)
        en_titles = re.findall(r"\"([^\"]+)\"", sent)
        year_m = re.search(r"(19|20)\d{2}", sent)
        year = year_m.group(0) if year_m else ""

        if "Zhang Wei" in sent and "Wang Xiaoming" in sent and len(cn_titles) >= 2 and year:
            items.append(
                f"Zhang Wei, Wang Xiaoming. ({year}). \"{cn_titles[0]}\". {cn_titles[1]}."
            )
            continue

        # Pair year with nearby titles when possible
        pairs = re.findall("((?:19|20)\\d{2})year[^《\\\"]{0,20}(《[^》]+》|\\\"[^\\\"]+\\\")", sent)
        if pairs:
            for y, t in pairs:
                items.append(f"{t}，{y}Year")
            if "paper" in sent and cn_titles and year:
                if year:
                    items.append(f"《{cn_titles[0]}》，{year}Year")
        else:
            if cn_titles:
                if len(cn_titles) >= 2:
                    title = f"《{cn_titles[0]}》，{year}Year，《{cn_titles[1]}》" if year else f"《{cn_titles[0]}》，《{cn_titles[1]}》"
                    items.append(title)
                else:
                    title = f"《{cn_titles[0]}》，{year}Year" if year else f"《{cn_titles[0]}》"
                    items.append(title)
            for t in en_titles:
                title = f"\"{t}\", {year}" if year else f"\"{t}\""
                items.append(title)
    return dedupe_keep_order(items)


def extract_narrative_projects(sentences):
    items = []
    for sent in sentences:
        if not any(k in sent for k in ("project", "platform", "system", "develop", "design")):
            continue
        clauses = [c.strip("，。 ") for c in re.split("(?: at the same time | in addition | and | and)", sent) if c.strip()]
        for clause in clauses:
            if not any(k in clause for k in ("project", "platform", "system")):
                continue
            name_m = re.search("(?:Development|Participate|Responsible|Design).{0,20}?([\\u4e00-\\u9fffA-Za-z0-9]+(?:Platform|System|Project))", clause)
            if not name_m:
                name_m = re.search("([\\u4e00-\\u9fffA-Za-z0-9]+(?:Platform|System|Project))", clause)
            if not name_m:
                continue
            name = name_m.group(1)
            name = re.sub("^(He was responsible for developing one | He also participated in one | participated in one | was responsible for developing one | one | one)", "", name).strip()
            name_parts = re.findall("[\\u4e00-\\u9fffA-Za-z0-9]+(?:Platform|System|Project)", name)
            if name_parts:
                name = name_parts[-1]
            time = ""
            m = re.search("(19|20)\\d{2} year\\d{1,2} month.*?(19|20)\\d{2} year\\d{1,2} month", clause)
            if m:
                time = m.group(0)
            else:
                m = re.search("From (19|20)\\d{2} year\\d{1,2} month.*?(19|20)\\d{2} year\\d{1,2} month", clause)
                if m:
                    time = m.group(0).replace("from", "")
            time = time.replace("to", " - ").replace("arrive", " - ").replace("last until", " - ")
            time = time.replace("continued", "").replace("  ", " ").strip()
            desc = clause
            desc = re.sub("^.*?(Responsible for | Participate in | Development | Design)", "", desc, count=1)
            if name in desc:
                desc = desc.split(name, 1)[-1]
            if time:
                desc = desc.replace(time, "")
            desc = re.sub("This project was launched in [^,. ]*[,. ]", "", desc)
            desc = re.sub("This project starts from [^,. ]*[,. ]", "", desc)
            desc = desc.replace("The design and implementation of", "")
            desc = desc.strip("，。 ")
            desc = re.sub("^(aim to|by|utilize|and|and|while|with)", "", desc)
            title = f"{name}（{time}）" if time else name
            items.append(f"{title}：{desc}" if desc else title)
    return dedupe_keep_order(items)


def extract_narrative_awards(sentences):
    items = []
    for sent in sentences:
        if not any(k in sent for k in ("award", "funding", "fund")):
            continue
        parts = re.split("[,;;]|and|including", sent)
        for part in parts:
            part = part.strip("，。 ")
            m = re.search("((19|20)\\d{2})year[^,. ]*?(Award | Funding | Fund)", part)
            if m:
                year = f"{m.group(1)}Year"
                rest = part.split(year, 1)[-1].strip("，。 ")
                rest = re.sub("^(of|obtained|obtained|obtained|in|for)", "", rest)
                rest = rest.lstrip("of")
                items.append(f"{year}，{rest}" if rest else year)
    return dedupe_keep_order(items)


def extract_narrative_skills(text):
    items = []
    lang_tokens = ["Python", "Java", "C++", "SQL", "NoSQL", "R", "MATLAB", "Go", "Rust", "C#"]
    domain_tokens = ["machine learning", "deep learning", "data mining", "AI", "computer vision", "biological information"]
    tool_tokens = ["TensorFlow", "PyTorch", "SQL", "NoSQL"]
    langs = [t for t in lang_tokens if t in text]
    domains = [t for t in domain_tokens if t in text]
    tools = [t for t in tool_tokens if t in text]
    if langs:
        items.append(f"programming language：{', '.join(dedupe_keep_order(langs))}")
    if domains:
        items.append(f"Technical field：{', '.join(dedupe_keep_order(domains))}")
    if tools:
        items.append(f"Tools & Techniques：{', '.join(dedupe_keep_order(tools))}")
    return items


def infer_skills_from_projects(projects):
    text = " ".join(projects)
    skills = []
    # Heuristic mapping from common project domains to skills
    if any(k in text for k in ("Blockchain", "chain")):
        skills.append("Blockchain")
    if any(k in text for k in ("vaccine", "biology", "Gene", "cell")):
        skills.append("bioinformatics")
    if any(k in text for k in ("medical", "diagnosis", "patient")):
        skills.append("Medical data analysis")
    if any(k in text for k in ("machine learning", "deep learning", "AI", "AI")):
        skills.append("machine learning")
    if any(k in text for k in ("data analysis", "data mining", "predict")):
        skills.append("data analysis")
    if any(k in text for k in ("system", "platform")):
        skills.append("system development")
    if not skills:
        return []
    return [f"Technical field：{', '.join(dedupe_keep_order(skills))}"]


def normalize_project_items(lines):
    items = []
    current = {"name": "", "time": "", "desc": []}

    def parse_year_project(line):
        m = re.match("(?P<year>(19|20)\\d{2}year)[::]\\s*(?P<body>.+)", line)
        if not m:
            return ""
        year = m.group("year")
        body = m.group("body")
        name_m = re.search(r"“([^”]+)”", body)
        name = name_m.group(1) if name_m else ""
        role_m = re.search(r"（([^）]+)）", body)
        role = role_m.group(1) if role_m else ""
        action_m = re.search("(Responsible | Participate | Host)", body)
        action = action_m.group(1) if action_m else ""
        if action and role and action == role:
            role = ""
        if not name:
            name_m = re.search("([\\u4e00-\\u9fffA-Za-z0-9]+(?:Project|System|Program|Platform|Research))", body)
            name = name_m.group(1) if name_m else body
        desc_parts = [p for p in [action, role] if p]
        desc = "，".join(desc_parts) if desc_parts else ""
        return f"{name}（{year}）：{desc}" if desc else f"{name}（{year}）"

    def flush():
        if not current["name"] and not current["desc"]:
            return
        name = current["name"] or ""
        time = current["time"]
        desc = " ".join(current["desc"]).strip()
        if time:
            title = f"{name}（{time}）" if name else f"（{time}）"
        else:
            title = name
        if desc:
            items.append(f"{title}：{desc}" if title else desc)
        else:
            if title:
                items.append(title)

    for ln in lines:
        ln = polish_line(strip_leading_markers(ln))
        parsed = parse_year_project(ln)
        if parsed:
            flush()
            items.append(parsed)
            current = {"name": "", "time": "", "desc": []}
            continue
        if ln.startswith("Project name"):
            flush()
            current = {"name": ln.split("：")[-1].split(":")[-1].strip(), "time": "", "desc": []}
            continue
        if ln.startswith("time") or ln.lower().startswith("time"):
            current["time"] = ln.split("：")[-1].split(":")[-1].strip()
            continue
        if ln.startswith("describe") or ln.lower().startswith("description"):
            desc = ln.split("：")[-1].split(":")[-1].strip()
            if desc:
                current["desc"].append(desc)
            continue
        if re.match("\\d{4}year", ln):
            current["time"] = ln
            continue
        current["desc"].append(ln)
    flush()
    return items


def normalize_skill_items(lines):
    items = []
    for ln in lines:
        ln = polish_line(strip_leading_markers(ln))
        if ln == "Skill" or ln.lower() == "skills":
            continue
        items.append(ln)
    return items


def extract_contact(lines):
    name = ""
    title = ""
    org = ""
    email = ""
    phone = ""
    for ln in lines[:10]:
        ln = strip_leading_markers(ln)
        if not name and ("Name" in ln or ln.lower().startswith("name")):
            name = ln.split("：")[-1].split(":")[-1].strip()
        elif not title and ("Position" in ln or "job title" in ln):
            title = ln.split("：")[-1].split(":")[-1].strip()
        elif not org and ("School" in ln or "unit" in ln or "mechanism" in ln or "Institution" in ln):
            org = ln.split("：")[-1].split(":")[-1].strip()
        if not email:
            m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", ln)
            if m:
                email = m.group(0)
        if not phone:
            m = re.search(r"(\+?\d[\d\- ]{7,}\d)", ln)
            if m:
                phone = m.group(1).strip()

        if not title and re.search("(Assistant Professor | Associate Professor | Professor | Lecturer)", ln):
            title = re.search("(Assistant Professor | Associate Professor | Professor | Lecturer)", ln).group(1)
        if not org and re.search("([\\u4e00-\\u9fff]{2,30}(University|College|Institute|Department))", ln):
            org = re.search("([\\u4e00-\\u9fff]{2,30}(University|College|Institute|Department))", ln).group(1)
        if not name:
            if re.fullmatch(r"[\u4e00-\u9fff]{2,4}", ln):
                name = ln
            elif re.fullmatch(r"[A-Za-z][A-Za-z .'-]{1,}", ln):
                name = ln
    return name, title, org, email, phone


def extract_contact_from_text(text):
    name = ""
    title = ""
    org = ""
    email = ""
    phone = ""
    m = re.search("([\\u4e00-\\u9fff]{2,4}) is (one bit)?", text)
    if m:
        name = m.group(1)
    if not name:
        m = re.search("Named ([\\u4e00-\\u9fff]{2,4})", text)
        if m:
            name = m.group(1)
    m = re.search("at ([\\u4e00-\\u9fff]{2,30}(University|College|Institute|Department))", text)
    if m:
        org = m.group(1)
    if not org:
        m = re.search("([\\u4e00-\\u9fff]{2,30}(University|College|Institute|Department))", text)
        if m:
            org = re.sub("^(currently in|in)", "", m.group(1))
    m = re.search("(Assistant Professor | Associate Professor | Professor | Lecturer)", text)
    if m:
        title = m.group(1)
    m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    if m:
        email = m.group(0)
    m = re.search(r"(\+?\d[\d\- ]{7,}\d)", text)
    if m:
        phone = m.group(1).strip()
    return name, title, org, email, phone


def update_contact_from_line(line, name, title, org, email, phone):
    if line.startswith("Name") and not name:
        name = line.split("：")[-1].split(":")[-1].strip()
    if line.startswith(("Position", "job title")) and not title:
        title = line.split("：")[-1].split(":")[-1].strip()
    if line.startswith(("School", "unit", "mechanism", "Institution")) and not org:
        org = line.split("：")[-1].split(":")[-1].strip()
    if ("Mail" in line or "Email" in line) and not email:
        m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", line)
        if m:
            email = m.group(0)
    if "Telephone" in line and not phone:
        m = re.search(r"(\+?\d[\d\- ]{7,}\d)", line)
        if m:
            phone = m.group(1).strip()
    return name, title, org, email, phone


def parse_text(text):
    lines = clean_lines(text)
    narrative_mode = len(lines) <= 2 or not has_explicit_headers(lines)
    if narrative_mode:
        name, title, org, email, phone = extract_contact_from_text(text)
    else:
        name, title, org, email, phone = extract_contact(lines)

    sections = {k: [] for k in HEADER_MAP.keys()}
    summary = ""
    current = ""
    work_lines = lines if not narrative_mode else split_sentences(text)
    contact_values = {v for v in (name, title, org, email, phone) if v}

    for ln in work_lines:
        ln = polish_line(strip_leading_markers(ln))
        # Always try to update contact info when explicit fields appear
        if any(ln.startswith(k) for k in ("Name", "Position", "job title", "School", "unit", "mechanism", "Institution", "Mail", "Email", "Telephone")):
            name, title, org, email, phone = update_contact_from_line(
                ln, name, title, org, email, phone
            )
            contact_values = {v for v in (name, title, org, email, phone) if v}
            continue

        if ln in contact_values or is_contact_line(ln):
            continue

        header_key = normalize_header_key(ln)
        if header_key in CN_HEADERS + EN_HEADERS:
            current = classify_section(header_key, current, bool(summary))
            continue

        section = classify_section(ln, current, bool(summary))
        current = section

        if current == "profile":
            summary = ln
        elif current:
            sections[current].append(ln)

    if narrative_mode:
        sentences = split_sentences(text)
        sections["education"] = extract_narrative_education(sentences, org)
        sections["publications"] = extract_narrative_publications(sentences)
        sections["projects"] = extract_narrative_projects(sentences)
        sections["awards"] = extract_narrative_awards(sentences)
        sections["skills"] = extract_narrative_skills(text)

    if not narrative_mode:
        sections["projects"] = normalize_project_items(sections["projects"])
    sections["skills"] = normalize_skill_items(sections["skills"])
    sections = {k: [polish_line(it) for it in v] for k, v in sections.items()}

    return {
        "name": name or "Academic CV",
        "title": title,
        "org": org,
        "email": email,
        "phone": phone,
        "sections": sections,
    }


def parse_education_item(item):
    if "：" in item:
        item = item.replace("：", "，", 1)
    degree = ""
    for d in ("PhD", "master", "Bachelor"):
        if d in item:
            degree = d
            break
    parts = [p.strip() for p in re.split(r"，|,", item) if p.strip()]
    major = parts[1] if len(parts) > 1 else ""
    school = parts[2] if len(parts) > 2 else ""
    year = ""
    for p in parts:
        if re.search(r"(19|20)\d{2}", p):
            year = p
            break
    if year:
        parts = [p for p in parts if p != year]
    return degree, major, school, year


def sort_education(items):
    def key(it):
        degree, _, _, year = parse_education_item(it)
        rank = DEGREE_RANK.get(degree, 0)
        year_n = int(re.search(r"(19|20)\d{2}", year).group(0)) if year else 0
        return (-rank, -year_n)

    return sorted(items, key=key)


def extract_end_year(item):
    years = re.findall(r"(19|20)\d{2}", item)
    if not years:
        return 0
    return int(years[-1])


def sort_by_year_desc(items):
    return sorted(items, key=extract_end_year, reverse=True)


def build_markdown(data):
    name = data.get("name", "")
    title = data.get("title", "")
    org = data.get("org", "")
    email = data.get("email", "")
    phone = data.get("phone", "")

    education = sort_education(data["sections"].get("education", []))
    projects = sort_by_year_desc(data["sections"].get("projects", []))
    publications = data["sections"].get("publications", [])
    awards = sort_by_year_desc(data["sections"].get("awards", []))
    skills = data["sections"].get("skills", [])
    if not skills:
        skills = infer_skills_from_projects(projects)

    def format_education(item):
        degree, major, school, year = parse_education_item(item)
        parts = [p for p in (degree, major, school, year) if p]
        return "，".join(parts)

    md = []
    md.append("personal data")
    md.append(" | ".join([name, title, org, email, phone]))
    md.append("")
    md.append("Educational background")
    for it in education:
        md.append(format_education(it))
    md.append("")
    md.append("Project experience")
    for it in projects:
        md.append(it)
    md.append("")
    md.append("Paper published")
    for it in publications:
        md.append(it)
    md.append("")
    if awards:
        md.append("Awards")
        for it in awards:
            md.append(it)
        md.append("")
    if skills:
        md.append("Skill")
        for it in skills:
            md.append(it)

    return "\n".join(md).strip() + "\n"


def render_docx_from_markdown(md_text, color, labels):
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    lines = [ln.rstrip() for ln in md_text.splitlines()]
    current_section = ""
    profile_name = ""
    profile_contact = ""
    for ln in lines:
        if not ln.strip():
            continue
        key = normalize_header_key(ln)
        if key in ("personal data", "Educational background", "Project experience", "Paper published", "Awards", "Skill"):
            current_section = key
            if key != "personal data":
                p = doc.add_paragraph()
                run = p.add_run(key)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = color
            continue

        if current_section == "personal data" and "|" in ln:
            parts = [p.strip() for p in ln.split("|")]
            profile_name = parts[0] if len(parts) > 0 else ""
            title = parts[1] if len(parts) > 1 else ""
            org = parts[2] if len(parts) > 2 else ""
            email = parts[3] if len(parts) > 3 else ""
            phone = parts[4] if len(parts) > 4 else ""
            bits = [bit for bit in (title, org, email, phone) if bit]
            profile_contact = " | ".join(bits)
            continue

        if current_section in ("Educational background", "Project experience", "Paper published", "Awards", "Skill"):
            doc.add_paragraph(ln.lstrip("• "), style="List Bullet")
        else:
            doc.add_paragraph(ln.lstrip("• "))

    if profile_name and doc.paragraphs:
        first = doc.paragraphs[0]
        title_p = first.insert_paragraph_before()
        run = title_p.add_run(profile_name)
        run.bold = True
        run.font.size = Pt(18)
        if profile_contact:
            contact_p = first.insert_paragraph_before(profile_contact)
            contact_p.runs[0].font.size = Pt(10)

    return doc


def main():
    parser = argparse.ArgumentParser(description="Render academic CV DOCX from free text.")
    parser.add_argument("--input", required=True, help="Input text file")
    parser.add_argument("--output", required=True, help="Output DOCX path or directory")
    parser.add_argument("--lang", choices=["zh", "en"], help="Force output language")
    parser.add_argument(
        "--header-color",
        choices=["purple", "cyan", "green", "red", "random"],
        default="random",
        help="Header color (unified color)",
    )
    args = parser.parse_args()

    text = Path(args.input).read_text(encoding="utf-8")
    lang = args.lang or detect_lang(text)
    data = parse_text(text)

    color_key = args.header_color
    if color_key == "random":
        color_key = random.choice(["purple", "cyan", "green", "red"])
    color = COLOR_MAP[color_key]

    md_text = build_markdown(data)
    doc = render_docx_from_markdown(md_text, color, HEADER_MAP)

    out_path = Path(args.output)
    target_name = f"{data['name']}-academic resume.docx"
    if out_path.is_dir():
        out_path = out_path / target_name
    else:
        out_path = out_path.parent / target_name
    doc.save(out_path)
    print(f"DOCX saved to {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
