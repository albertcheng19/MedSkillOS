#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Text formatting module - Word format reserved version
Clean blank lines and spaces while preserving Word document formatting"""

import re
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import List, Optional, Tuple, Dict


class LineEnding(Enum):
    UNIX = "\n"
    WINDOWS = "\r\n"
    MAC = "\r"


class IndentType(Enum):
    SPACES = "spaces"
    TABS = "tabs"


@dataclass
class FormatOptions:
    remove_empty_lines: bool = True
    max_empty_lines: int = 2
    trim_line_end_spaces: bool = True
    normalize_tabs: bool = True
    tab_to_spaces: bool = True
    indent_type: IndentType = IndentType.SPACES
    indent_size: int = 4
    line_ending: LineEnding = LineEnding.UNIX
    preserve_formatting: bool = True
    remove_trailing_newline: bool = False


class TextFormatter:
    def __init__(self, options: FormatOptions = None):
        self.options = options or FormatOptions()

    def format(self, text: str) -> str:
        if not text:
            return text

        lines = text.splitlines(keepends=True)
        result_lines = []

        for line in lines:
            if self.options.trim_line_end_spaces:
                line = line.rstrip()
            result_lines.append(line)

        text = "".join(result_lines)

        if self.options.remove_empty_lines:
            max_lines = self.options.max_empty_lines
            pattern = rf"\n{{{max_lines},}}"
            replacement = f"\n{'' * max_lines}\n"
            text = re.sub(pattern, replacement, text)

        if self.options.line_ending == LineEnding.WINDOWS:
            text = text.replace("\n", "\r\n")
        elif self.options.line_ending == LineEnding.MAC:
            text = text.replace("\r\n", "\n").replace("\n", "\r")
        elif self.options.line_ending == LineEnding.UNIX:
            text = text.replace("\r\n", "\n").replace("\r", "\n")

        if self.options.remove_trailing_newline:
            text = text.rstrip("\n\r")

        return text

    def format_file(self, input_path: str, output_path: str = None) -> str:
        input_file = Path(input_path)
        if not input_file.exists():
            raise FileNotFoundError(f"File does not exist: {input_path}")

        text = input_file.read_text(encoding="utf-8")
        formatted = self.format(text)

        if output_path is None:
            output_path = str(
                input_file.parent / f"{input_file.stem}_clean{input_file.suffix}"
            )

        output_file = Path(output_path)
        output_file.write_text(formatted, encoding="utf-8")
        return str(output_file)

    def preview(self, text: str) -> str:
        return self.format(text)


def clean_docx_inplace(
    input_path: str,
    font_name: str = "Times New Roman",
    font_size: int = 12,
    remove_empty_lines: bool = True,
) -> dict:
    """Modify Word documents directly (preserving formatting)"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document(input_path)

        paragraphs_to_remove = []
        removed_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()

            is_empty = not text or text in ["()", "（ ）", "（ ）", "( )", "Â", "Ã"]

            if remove_empty_lines and is_empty:
                if para._element.getparent() is not None:
                    para._element.getparent().remove(para._element)
                    removed_count += 1
                continue

            if para.runs:
                for run in para.runs:
                    if run.font.name:
                        run.font.name = font_name
                    if run.font.size:
                        run.font.size = Pt(font_size)
                    run.font._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

            if para.style and para.style.font:
                if para.style.font.name:
                    para.style.font.name = font_name
                if para.style.font.size:
                    para.style.font.size = Pt(font_size)
                para.style.font._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        doc.save(input_path)

        return {
            "success": True,
            "file": input_path,
            "removed_empty": removed_count,
        }

    except ImportError:
        raise ImportError("Need to install python-docx: pip install python-docx")


def clean_docx_preserve_format(
    input_path: str,
    output_path: str = None,
    font_name: str = "Times New Roman",
    font_size: int = 12,
    remove_empty_lines: bool = True,
) -> str:
    """Clean Word document (make a copy) - retain all original formatting
    Use direct modification to retain all styles including EndNote Bibliography"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document(input_path)

        paragraphs_to_remove = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            empty_patterns = [
                "",
                "()",
                "（ ）",
                "（ ）",
                "( )",
                "Â",
                "Ã",
                "Â ",
                " Ã",
                "Â",
                "Ã",
                "â€˜",
                "â€",
                " ",
                " ",
                " ",
                " ",
                " ",
                "•",
                "·",
                "__________________",
                "_________",
                "_____",
                "--------------",
                "-------",
                "----",
            ]

            is_empty = (
                not text or text in empty_patterns or len(text.strip("._- ")) == 0
            )

            if remove_empty_lines and is_empty:
                paragraphs_to_remove.append(para)

        for para in paragraphs_to_remove:
            para._element.getparent().remove(para._element)

        if output_path is None:
            input_file = Path(input_path)
            output_path = str(
                input_file.parent / f"{input_file.stem}_clean{input_file.suffix}"
            )

        doc.save(output_path)
        return output_path

    except ImportError:
        raise ImportError("Need to install python-docx: pip install python-docx")


def extract_text_from_docx(docx_path: str) -> tuple:
    """Extract plain text and paragraph style information from Word documents"""
    try:
        from docx import Document

        doc = Document(docx_path)
        paragraphs = []
        para_styles = []

        for para in doc.paragraphs:
            text = para.text
            style_name = para.style.name if para.style else "Normal"
            paragraphs.append(text)
            para_styles.append(style_name)

        return "\n".join(paragraphs), para_styles

    except ImportError:
        raise ImportError("Need to install python-docx: pip install python-docx")
